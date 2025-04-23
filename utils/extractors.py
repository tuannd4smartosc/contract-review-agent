from docx import Document
from PyPDF2 import PdfReader
import re
import html

def clean_html_text(raw_html):
    text = re.sub(r'<.*?>', '', raw_html)
    text = html.unescape(text)
    text = re.sub(r'\s*\n\s*', '\n', text.strip())
    text = re.sub(r'\s{2,}', ' ', text)
    return text

def extract_text_from_pdf(pdf_path):
    text = ""
    try:
        with open(pdf_path, 'rb') as pdf_file:
            reader = PdfReader(pdf_file)
            for page_num in range(len(reader.pages)):
                page = reader.pages[page_num]
                text += page.extract_text()
        return text
    except Exception as e:
        return f"Error extracting text with PyPDF2: {e}"

def extract_text_from_docx(docx_path):
    doc = Document(docx_path)
    text = ""
    for para in doc.paragraphs:
        text += para.text + "\n"  
    return text

def extract_text_from_docx_with_tables(docx_path):
    doc = Document(docx_path)
    text = ""

    for para in doc.paragraphs:
        text += para.text + "\n"
    
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text += cell.text + " " 
            text += "\n"  

    return text